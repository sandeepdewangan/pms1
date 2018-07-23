from docx import Document
from docx.shared import Inches

document = Document()

print(dir(document))

document.add_heading('Index')

# get table data -------------
items = {'qty':10, 'sku':'ye kya hota hai', 'desc': 'Samsung AC'}

# add table ------------------
table = document.add_table(1, 3)

# populate header row --------
heading_cells = table.rows[0].cells
heading_cells[0].text = 'Qty'
heading_cells[1].text = 'SKU'
heading_cells[2].text = 'Description'

# add a data row for each item
for item in items:
    cells = table.add_row().cells
    cells[0].text = '1'
    cells[1].text = 'Kya hai Ye'
    cells[2].text = 'Samsung AC'


document.save('demo.docx')
